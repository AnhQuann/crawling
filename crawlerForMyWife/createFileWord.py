from urllib.request import urlopen
from bs4 import BeautifulSoup
from docx import Document
import html2text

document = Document()


def createFileWord(rootUrl, pathUrl, heading):
  htmlContent = urlopen(rootUrl + pathUrl).read().decode("utf8")
  soup = BeautifulSoup(htmlContent, 'html.parser')
  contentArea = soup.find("div", "col-md-9")

  converter = html2text.HTML2Text()
  converter.ignore_links = True
  converter.ignore_images = True
  converter.escape_all = True

  textContent = converter.handle(str(contentArea))

  document.add_heading(heading, 1)
  document.add_paragraph(textContent)


def saveFileWord(fileName):
  document.save(fileName)
